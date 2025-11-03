#!/usr/bin/env python3
"""
Generate standard manuscript format .docx from fiction project scenes.
Follows William Shunn's modern manuscript format specification.
"""

import os
import json
import re
import unicodedata
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def clean_invisible_unicode(text):
    """Remove invisible Unicode characters that AI might add."""
    # Common invisible characters
    invisible_chars = [
        '\u200B',  # Zero-width space
        '\u200C',  # Zero-width non-joiner
        '\u200D',  # Zero-width joiner
        '\u00AD',  # Soft hyphen
        '\uFEFF',  # Zero-width no-break space
        '\u2060',  # Word joiner
    ]
    for char in invisible_chars:
        text = text.replace(char, '')

    # Normalize Unicode (NFC form)
    text = unicodedata.normalize('NFC', text)

    return text

def parse_markdown_formatting(text):
    """
    Convert markdown to formatting tuples.
    Returns list of (text, is_italic, is_bold) tuples.
    """
    result = []

    # Pattern: **bold** or *italic*
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|___(.+?)___|__(.+?)__|_(.+?)_)'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before match (plain)
        if match.start() > last_end:
            plain_text = text[last_end:match.start()]
            if plain_text:
                result.append((plain_text, False, False))

        # Determine formatting
        full_match = match.group(0)
        inner_text = match.group(2) or match.group(3) or match.group(4) or \
                     match.group(5) or match.group(6) or match.group(7)

        if full_match.startswith('***') or full_match.startswith('___'):
            # Bold + Italic
            result.append((inner_text, True, True))
        elif full_match.startswith('**') or full_match.startswith('__'):
            # Bold only
            result.append((inner_text, False, True))
        elif full_match.startswith('*') or full_match.startswith('_'):
            # Italic only
            result.append((inner_text, True, False))

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            result.append((remaining, False, False))

    return result if result else [(text, False, False)]

def add_formatted_text(paragraph, text):
    """Add text to paragraph with markdown formatting preserved."""
    segments = parse_markdown_formatting(text)

    for segment_text, is_italic, is_bold in segments:
        run = paragraph.add_run(segment_text)
        run.italic = is_italic
        run.bold = is_bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def get_title_keywords(title):
    """Extract 1-2 keywords from title for header."""
    # Remove common words
    stopwords = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with'}
    words = [w for w in title.lower().split() if w not in stopwords]

    # Return first 2 significant words, or first word if short title
    if len(words) >= 2:
        return ' '.join(words[:2]).title()
    elif words:
        return words[0].title()
    else:
        return title.split()[0] if title.split() else 'Story'

def create_header(section, author_last_name, title_keywords):
    """Create header with Surname / Keywords / Page# format per Shunn spec."""
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Set font first
    run = header_para.add_run(f"{author_last_name} / {title_keywords} / ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Add page number field
    run = header_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def create_manuscript(project_dir, output_path):
    """Main function to create manuscript per Shunn format."""

    # Read project.json
    with open(os.path.join(project_dir, 'project.json'), 'r') as f:
        project = json.load(f)

    title = project.get('projectName', project.get('title', 'Untitled'))
    author = project.get('author', 'Unknown Author')
    pen_name = project.get('penName', author)
    contact = project.get('contact', {})
    word_count = project.get('wordCount', 0)

    # Round word count to nearest 100 (or 500 for novellas)
    if word_count > 17500:  # Novella territory
        word_count_rounded = round(word_count / 500) * 500
    else:
        word_count_rounded = round(word_count / 100) * 100

    # Get author last name
    author_last_name = author.split()[-1] if ' ' in author else author

    # Get title keywords for header
    title_keywords = get_title_keywords(title)

    # Create document
    doc = Document()

    # Set up page margins (1 inch all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # === TITLE PAGE (Page 1) ===
    # Upper left: Contact info
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = para.add_run(author)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Contact info continuation
    if contact:
        address = contact.get('address', [])
        if isinstance(address, list):
            for line in address:
                para = doc.add_paragraph(str(line))
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                run = para.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        elif address:
            para = doc.add_paragraph(str(address))
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            run = para.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        phone = contact.get('phone', '')
        if phone:
            para = doc.add_paragraph(phone)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            run = para.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        email = contact.get('email', '')
        if email:
            para = doc.add_paragraph(email)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            run = para.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # Upper right: Word count (we'll do this manually by going back)
    # Actually, we need to insert it on the first paragraph
    first_para = doc.paragraphs[0]
    first_para.add_run('\t' * 5 + f"about {word_count_rounded:,} words")

    # Blank lines to push title 1/3 to 1/2 down page
    # At double-spacing, ~12 blank lines gets us about 1/3 down
    for _ in range(12):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Title (centered)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = para.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Double-space once and add byline
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = para.add_run(f"by {pen_name}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Double-space twice more, then start text
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # === READ AND PROCESS SCENES ===
    scenes_dir = os.path.join(project_dir, 'scenes')
    scene_files = sorted([f for f in os.listdir(scenes_dir)
                         if f.startswith('scene-') and f.endswith('.md')])

    first_scene = True

    for idx, scene_file in enumerate(scene_files, 1):
        scene_path = os.path.join(scenes_dir, scene_file)

        with open(scene_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Clean invisible Unicode
        content = clean_invisible_unicode(content)

        # Strip metadata header
        # Format: # Scene NNN + metadata fields + --- separator + prose
        lines = content.split('\n')
        prose_start = 0

        # Find where prose actually starts
        for i, line in enumerate(lines):
            stripped = line.strip()

            # Skip scene heading (# Scene XXX)
            if stripped.startswith('# Scene'):
                prose_start = i + 1
                continue

            # Skip blank lines
            if not stripped:
                continue

            # Skip metadata fields (**Field**: value)
            if stripped.startswith('**') and '**:' in stripped:
                prose_start = i + 1
                continue

            # If we hit the --- separator, skip it and start prose after
            if stripped == '---':
                prose_start = i + 1
                break

            # If we hit regular text (not metadata), this is where prose starts
            if stripped and not stripped.startswith('**'):
                prose_start = i
                break

        # Rejoin from where prose starts
        content = '\n'.join(lines[prose_start:]).strip()

        # For scenes after the first, add page break (scenes = chapters)
        if not first_scene:
            doc.add_page_break()

        first_scene = False

        # Process paragraphs
        paragraphs = content.split('\n\n')

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Check for scene break markers within content
            if para_text in ['***', '* * *', '# # #', '#']:
                # Center scene break
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                run = para.add_run('#')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                continue

            # Regular paragraph
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            para.paragraph_format.space_after = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-aligned, ragged right

            # Add formatted text
            add_formatted_text(para, para_text)

    # === END MATTER ===
    # Optional: Add "END" centered
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = para.add_run("END")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Now add header to pages 2+ (not page 1)
    # This needs to be done by creating a new section after page 1
    # For simplicity, we'll apply header to all pages and rely on "different first page" setting

    # Enable different first page
    doc.sections[0].different_first_page_header_footer = True

    # Add header to subsequent pages (section 0, but not first page)
    create_header(doc.sections[0], author_last_name, title_keywords)

    # Save .docx
    doc.save(output_path)
    print(f"✓ Created: {output_path}")

    return output_path

if __name__ == '__main__':
    import sys

    project_dir = sys.argv[1] if len(sys.argv) > 1 else '.'
    output_file = sys.argv[2] if len(sys.argv) > 2 else 'manuscript.docx'

    create_manuscript(project_dir, output_file)
